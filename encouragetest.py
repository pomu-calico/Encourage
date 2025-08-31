from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 新しいWord文書を作成
doc = Document()

# ページ設定（A4縦、余白小さめ）
sections = doc.sections
for section in sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# タイトル
title = doc.add_paragraph("【第1問】次の文章を読んで、後の問いに答えよ。")
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
title.runs[0].font.size = Pt(14)
title.runs[0].bold = True

doc.add_paragraph("")

# 本文
p = doc.add_paragraph()
run = p.add_run("　ここに問題文が入ります。共通テスト風に長文を載せたい場合は、このように段落を使います。")
run.font.size = Pt(11)

doc.add_paragraph("")

# 設問
q = doc.add_paragraph("問1　この文章の内容として最も適当なものを、次の①～④のうちから一つ選べ。")
q.runs[0].font.size = Pt(11)
q.runs[0].bold = True

doc.add_paragraph("　①　選択肢A").runs[0].font.size = Pt(11)
doc.add_paragraph("　②　選択肢B").runs[0].font.size = Pt(11)
doc.add_paragraph("　③　選択肢C").runs[0].font.size = Pt(11)
doc.add_paragraph("　④　選択肢D").runs[0].font.size = Pt(11)

doc.add_paragraph("")
doc.add_paragraph("──────────────────────────────")

# 保存
doc.save("共通テスト風テンプレート.docx")
